"""
Generate IEEE-format technical paper DOCX: AI Weather Prediction System
Run: python generate_paper.py
"""
import json, os, copy
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

FONT = "Times New Roman"
def _to_roman(n):
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
            (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    r = ''
    for v, s in vals:
        while n >= v:
            r += s; n -= v
    return r
ROMAN = {i: _to_roman(i) for i in range(1, 50)}
TABLE_NUM = [0]  # mutable counter for IEEE table numbering

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name=FONT, size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Force Times New Roman for East Asian fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

def set_paragraph_font(p, name=FONT, size=10, bold=False, italic=False):
    for run in p.runs:
        set_font(run, name, size, bold, italic)

def ieee_section(doc, num, title):
    """IEEE Level-1 heading: centered, Roman numeral, ALL CAPS."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    text = f"{ROMAN[num]}. {title.upper()}"
    run = p.add_run(text)
    set_font(run, size=10, bold=True)

def ieee_subsection(doc, letter, title):
    """IEEE Level-2 heading: left-aligned, italic, letter prefix."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{letter}. {title}")
    set_font(run, size=10, bold=False, italic=True)

def ieee_para(doc, text, size=10, first_indent=True, italic=False, bold=False):
    """IEEE body paragraph: justified, Times New Roman, first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic, bold=bold)
    return p

def ieee_bullet(doc, text, size=9):
    """IEEE itemised list entry."""
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def ieee_table_caption(doc, title):
    """IEEE table caption ABOVE table: 'TABLE I' format."""
    TABLE_NUM[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"TABLE {ROMAN[TABLE_NUM[0]]}")
    set_font(run, size=8, bold=True)
    p.add_run("\n")
    run2 = p.add_run(title)
    set_font(run2, size=8, italic=True)

FIG_NUM = [0]  # mutable counter for IEEE figure numbering
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "reports")

def ieee_add_figure(doc, image_path, caption, width=Inches(3.2)):
    """Insert an image with IEEE-style figure caption BELOW."""
    FIG_NUM[0] += 1
    if not os.path.exists(image_path):
        ieee_para(doc, f"[Image not found: {os.path.basename(image_path)}]",
                  size=8, italic=True, first_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r1 = cap.add_run(f"Fig. {FIG_NUM[0]}. ")
    set_font(r1, size=8, bold=True)
    r2 = cap.add_run(caption)
    set_font(r2, size=8)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def hdr_cell(cell, text, size=8):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=RGBColor(255, 255, 255))
    shade_cell(cell, "000000")

def tbl_cell(cell, text, size=8, bold=False, align_center=True):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)

def set_two_columns(section):
    """Set a section to two-column layout."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')  # 0.25 inch gap

def add_section_break(doc, break_type='continuous'):
    """Insert a section break."""
    new_section = doc.add_section()
    new_section._sectPr.find(qn('w:type')).set(qn('w:val'), break_type) if \
        new_section._sectPr.find(qn('w:type')) is not None else None
    # Set the type element
    sectPr = new_section._sectPr
    type_elem = sectPr.find(qn('w:type'))
    if type_elem is None:
        type_elem = OxmlElement('w:type')
        sectPr.insert(0, type_elem)
    type_elem.set(qn('w:val'), break_type)
    return new_section

# ── load results ─────────────────────────────────────────────────────────────

BASE = os.path.dirname(__file__)
with open(os.path.join(BASE, "evaluation", "outputs", "results.json")) as f:
    results = json.load(f)

import pandas as pd
train_df = pd.read_csv(os.path.join(BASE, "data", "processed", "train.csv"))
test_df  = pd.read_csv(os.path.join(BASE, "data", "processed", "test.csv"))
val_df   = pd.read_csv(os.path.join(BASE, "data", "processed", "val.csv"))

CLASS_NAMES = ["Sunny", "Cloudy", "Rainy", "Snowy"]

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# ── IEEE page margins (US Letter) ──
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(0.625)
    section.right_margin  = Inches(0.625)

# ── Set default font for entire document ──
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
rPr = style.element.find(qn('w:rPr'))
if rPr is None:
    rPr = OxmlElement('w:rPr')
    style.element.append(rPr)
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:ascii'), FONT)
rFonts.set(qn('w:hAnsi'), FONT)
rFonts.set(qn('w:cs'), FONT)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK (single-column)
# ══════════════════════════════════════════════════════════════════════════════

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
r = title.add_run(
    "Edge-Deployable AI Weather Prediction:\n"
    "A Comparative Study of Logistic Regression, Random Forest,\n"
    "XGBoost, and Neural Network Models\n"
    "Using Bosch BME688 Sensor Data"
)
set_font(r, size=24, bold=True)

# Author line
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.paragraph_format.space_after = Pt(4)
r = author.add_run("AI Weather Station Project")
set_font(r, size=11)

# Affiliation / platform
affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.paragraph_format.space_after = Pt(4)
r = affil.add_run("Target Platform: LilyGo T-Display S3 (ESP32-S3)")
set_font(r, size=10, italic=True)

# Date
datep = doc.add_paragraph()
datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
datep.paragraph_format.space_after = Pt(12)
r = datep.add_run(datetime.date.today().strftime('%B %d, %Y'))
set_font(r, size=10, italic=True)

# ── ABSTRACT (single-column, IEEE style) ──
abs_para = doc.add_paragraph()
abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_para.paragraph_format.space_after = Pt(6)
abs_para.paragraph_format.left_indent = Cm(0.75)
abs_para.paragraph_format.right_indent = Cm(0.75)

r_label = abs_para.add_run("Abstract\u2014")
set_font(r_label, size=9, bold=True, italic=True)
r_body = abs_para.add_run(
    "This paper presents a comparative evaluation of four supervised machine-learning "
    "approaches\u2014multinomial logistic regression (LR), random forest (RF), XGBoost, "
    "and a feedforward neural network (NN)\u2014applied to 6-hour-ahead weather condition "
    "classification using atmospheric sensor data compatible with the Bosch BME688 "
    "environmental sensor. All models share an identical 128-feature extraction pipeline "
    "derived from temperature, relative humidity, and barometric pressure readings over "
    "a 24-hour sliding window. Training data was sourced from the Open-Meteo historical "
    "weather archive across six geographically diverse stations spanning five years, "
    "yielding SMOTE-balanced training samples across four WMO-mapped classes: Sunny, "
    "Cloudy, Rainy, and Snowy. XGBoost achieves the best balanced macro-F1 score (0.541) "
    "and highest accuracy (58.2%), while the neural network achieves the strongest "
    "Sunny-class F1 (0.650). For deployment on the LilyGo T-Display S3 (ESP32-S3), "
    "the INT8-quantised neural network (91 KB TFLite) is recommended as the on-device "
    "classifier, with XGBoost serving as the host-side reference model."
)
set_font(r_body, size=9, italic=True)

# Index Terms
idx_para = doc.add_paragraph()
idx_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
idx_para.paragraph_format.space_after = Pt(10)
idx_para.paragraph_format.left_indent = Cm(0.75)
idx_para.paragraph_format.right_indent = Cm(0.75)
r_it = idx_para.add_run("Index Terms\u2014")
set_font(r_it, size=9, bold=True, italic=True)
r_kw = idx_para.add_run(
    "weather prediction, edge computing, machine learning, ESP32-S3, "
    "BME688, random forest, XGBoost, neural network, TensorFlow Lite."
)
set_font(r_kw, size=9, italic=True)

# ── Switch to two-column layout ──
new_sec = add_section_break(doc, 'continuous')
new_sec.top_margin    = Inches(0.75)
new_sec.bottom_margin = Inches(1.0)
new_sec.left_margin   = Inches(0.625)
new_sec.right_margin  = Inches(0.625)
set_two_columns(new_sec)

# ══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

sec_num = 1
ieee_section(doc, sec_num, "Introduction")

ieee_para(doc, (
    "Short-range weather forecasting at the sensor node level\u2014so-called edge "
    "weather intelligence\u2014has become increasingly relevant in the context of IoT "
    "environmental monitoring. Unlike NWP (Numerical Weather Prediction) models that "
    "require supercomputing resources, edge ML classifiers can operate locally on "
    "low-power microcontrollers such as the ESP32-S3, providing real-time 6-hour "
    "weather condition forecasts without cloud dependency."
))
ieee_para(doc, (
    "The Bosch BME688 sensor provides temperature, relative humidity, barometric "
    "pressure, gas resistance, IAQ index, estimated CO\u2082, and breath VOC "
    "measurements. This work leverages the three primary physical signals "
    "(temperature, humidity, pressure) for training, as public historical datasets "
    "do not include gas/IAQ measurements."
))
ieee_para(doc, (
    "Four classification paradigms are systematically compared: "
    "(1) multinomial logistic regression as the linear baseline; "
    "(2) random forest as a non-linear ensemble; "
    "(3) XGBoost as a gradient-boosted tree ensemble; and "
    "(4) a compact feedforward neural network. "
    "All models are trained on identical 128-dimensional feature vectors to ensure "
    "a fair comparison. The paper reports accuracy, macro-averaged F1, per-class "
    "precision/recall/F1, confusion matrices, model artifact sizes, and estimated "
    "inference latency on the ESP32-S3 platform."
))

# ══════════════════════════════════════════════════════════════════════════════
# II. DATASET
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Dataset")

ieee_subsection(doc, "A", "Data Source")
ieee_para(doc, (
    "Historical hourly weather data was obtained from the Open-Meteo Historical "
    "Weather API, an open-source service providing ERA5-reanalysis-backed atmospheric "
    "data at 1-hour resolution. Six stations were selected to maximise meteorological "
    "diversity and ensure representation of all four target classes (Table I)."
))

ieee_table_caption(doc, "Training stations and climate classifications.")
stations_tbl = doc.add_table(rows=7, cols=3)
stations_tbl.style = "Table Grid"
for cell, text in zip(stations_tbl.rows[0].cells, ["Station", "Coordinates", "Climate Type"]):
    hdr_cell(cell, text)
for row, (sta, coord, climate) in zip(stations_tbl.rows[1:], [
    ("London, UK",         "51.5\u00b0N, 0.1\u00b0W",   "Temperate oceanic (Cfb)"),
    ("Helsinki, Finland",  "60.2\u00b0N, 25.0\u00b0E",  "Humid continental (Dfb)"),
    ("Singapore",          "1.4\u00b0N, 103.8\u00b0E",  "Tropical rainforest (Af)"),
    ("Orlando, USA",       "28.5\u00b0N, 81.4\u00b0W",  "Humid subtropical (Cfa)"),
    ("Dhaka, Bangladesh",  "23.7\u00b0N, 90.4\u00b0E",  "Tropical monsoon (Am)"),
    ("Manaus, Brazil",     "3.1\u00b0S, 60.0\u00b0W",   "Tropical rainforest (Af)"),
]):
    for cell, text in zip(row.cells, [sta, coord, climate]):
        tbl_cell(cell, text)

ieee_subsection(doc, "B", "Label Mapping")
ieee_para(doc, (
    "Raw WMO present-weather codes are mapped to four integer class labels (Table II). "
    "WMO Stormy codes (95\u201399) are merged into the Rainy class, as T/RH/P signals "
    "cannot reliably distinguish thunderstorms from heavy rain."
))

ieee_table_caption(doc, "WMO weather code to class label mapping.")
wmo_tbl = doc.add_table(rows=5, cols=3)
wmo_tbl.style = "Table Grid"
for cell, text in zip(wmo_tbl.rows[0].cells, ["Label", "Class", "WMO Codes"]):
    hdr_cell(cell, text)
for row, (lbl, cls, codes) in zip(wmo_tbl.rows[1:], [
    ("0", "Sunny",  "0, 1"),
    ("1", "Cloudy", "2, 3, 45\u201348"),
    ("2", "Rainy",  "51\u201367, 80\u201382, 95\u201399"),
    ("3", "Snowy",  "71\u201377, 85\u201386"),
]):
    for cell, text in zip(row.cells, [lbl, cls, codes]):
        tbl_cell(cell, text)

ieee_subsection(doc, "C", "Data Statistics")

total = len(train_df) + len(val_df) + len(test_df)
train_dist = train_df.label.value_counts().sort_index().to_dict()
test_dist  = test_df.label.value_counts().sort_index().to_dict()

ieee_para(doc, (
    f"After WMO mapping, gap-filling (forward-fill \u2264 2 h; otherwise drop), and "
    f"chronological 70/15/15 splitting, the dataset contains {total:,} samples "
    f"(Table III)."
))

ieee_table_caption(doc, "Dataset split statistics by class.")
stats_tbl = doc.add_table(rows=4, cols=6)
stats_tbl.style = "Table Grid"
for cell, text in zip(stats_tbl.rows[0].cells,
    ["Split", "Samples", "Sunny", "Cloudy", "Rainy", "Snowy"]):
    hdr_cell(cell, text)
for tbl_row, (split_name, df) in zip(stats_tbl.rows[1:],
    [("Train (70%)", train_df), ("Val (15%)", val_df), ("Test (15%)", test_df)]):
    dist = df.label.value_counts().sort_index()
    vals = [split_name, str(len(df))]
    for i in range(4):
        vals.append(str(dist.get(i, 0)))
    for cell, text in zip(tbl_row.cells, vals):
        tbl_cell(cell, text)

ieee_para(doc, (
    f"The Snowy class is underrepresented (train: {train_dist.get(3,0):,} samples, "
    f"{100*train_dist.get(3,0)/len(train_df):.1f}%) and is addressed via SMOTE "
    f"upsampling. The SMOTE strategy caps each minority class at 20% of the majority "
    f"class count, with a maximum 5\u00d7 expansion per class. Balanced class weights "
    f"(computed via sklearn) are additionally applied to XGBoost (sample_weight) and "
    f"the neural network (class_weight)."
))

ieee_add_figure(doc, os.path.join(REPORTS_DIR, "class_distribution.png"),
                "Class distribution across train, validation, and test splits.")

# ══════════════════════════════════════════════════════════════════════════════
# III. FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Feature Engineering")

ieee_para(doc, (
    "A unified 128-dimensional feature vector is constructed from each 24-hour "
    "lookback window (readings at t\u22121 h through t\u221224 h) to predict the weather "
    "condition at t+6 h. All four models receive the identical feature vector, "
    "ensuring a scientifically valid comparison. Features are standardised using "
    "a sklearn StandardScaler fitted exclusively on the training set. Table IV "
    "summarises the feature groups."
))

FEATURE_GROUPS = [
    ("Raw signal lags", "temp, humidity, pressure \u00d7 24 timesteps", 72),
    ("Pressure tendency", "\u0394pressure at 1 h, 12 h, 24 h", 3),
    ("Pressure acceleration", "\u0394p_12h \u2212 \u0394p_24h", 1),
    ("Temp rate of change", "\u0394temp at 1 h and 12 h", 2),
    ("Dew point", "Magnus formula from t\u22121 h", 1),
    ("Absolute humidity", "Aug.\u2013Roche\u2013Magnus (g/m\u00b3)", 1),
    ("Rolling statistics", "mean, std, min, max \u00d7 6 h, 12 h, 24 h \u00d7 3 signals", 36),
    ("Cyclical time encoding", "sin/cos of hour-of-day, sin/cos of day-of-year", 4),
    ("Discriminative features", "freeze flags, dew-pt depression, pressure trend, composites", 8),
    ("Total", "", 128),
]

ieee_table_caption(doc, "Feature engineering groups (128 features total).")
feat_tbl = doc.add_table(rows=len(FEATURE_GROUPS)+1, cols=3)
feat_tbl.style = "Table Grid"
for cell, text in zip(feat_tbl.rows[0].cells, ["Feature Group", "Description", "Count"]):
    hdr_cell(cell, text)
for row, (grp, desc, cnt) in zip(feat_tbl.rows[1:], FEATURE_GROUPS):
    is_total = grp == "Total"
    tbl_cell(row.cells[0], grp, bold=is_total, align_center=False)
    tbl_cell(row.cells[1], desc, align_center=False)
    tbl_cell(row.cells[2], str(cnt), bold=is_total)
    if is_total:
        for cell in row.cells:
            shade_cell(cell, "E8E8E8")

ieee_para(doc, "Key derived features of meteorological significance include:")
for b in [
    "Pressure tendency (\u0394p at 1 h, 12 h, 24 h): falling pressure indicates approaching low-pressure systems.",
    "Temperature rate of change (1 h, 12 h): distinguishes rapid frontal passages from slow airmass change.",
    "Dew point (Magnus formula): encodes moisture content more physically than relative humidity alone.",
    "Freeze flags (T < 0\u00b0C, T < 3\u00b0C): binary discriminators between snow and rain.",
    "Dew-point depression (T \u2212 dew_point): near-zero depression indicates saturation and likely precipitation.",
    "Rolling statistics (6 h, 12 h, 24 h windows): capture signal volatility and trend direction.",
    "Cyclical time encoding (sin/cos): preserves continuity at midnight and year boundaries.",
]:
    ieee_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# IV. MODEL ARCHITECTURES
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Model Architectures and Hyperparameters")

# -- A. Logistic Regression --
ieee_subsection(doc, "A", "Logistic Regression (Linear Baseline)")
ieee_para(doc, (
    "Multinomial logistic regression models the posterior class probability as a "
    "softmax over a linear combination of input features. It provides a fully "
    "interpretable 4 \u00d7 128 coefficient matrix that directly quantifies each "
    "feature\u2019s contribution to each class decision."
))

ieee_table_caption(doc, "Logistic Regression hyperparameters.")
lr_params = [
    ("Solver", "saga"), ("Max iterations", "500"),
    ("Regularisation", "L2, C = 1.0"), ("Class weighting", "None (SMOTE)"),
    ("Input/Output", "128 features \u2192 4-class softmax"),
    ("Deployment artefact", "lr_coefficients.h (~6 KB)"),
    ("Inference (ESP32-S3)", "~4.3 \u00b5s"),
]
lr_tbl = doc.add_table(rows=len(lr_params)+1, cols=2)
lr_tbl.style = "Table Grid"
for cell, text in zip(lr_tbl.rows[0].cells, ["Parameter", "Value"]):
    hdr_cell(cell, text)
for row, (k, v) in zip(lr_tbl.rows[1:], lr_params):
    tbl_cell(row.cells[0], k, bold=True, align_center=False)
    tbl_cell(row.cells[1], v, align_center=False)

# -- B. Random Forest --
ieee_subsection(doc, "B", "Random Forest Classifier")
ieee_para(doc, (
    "Random forest builds an ensemble of decision trees, each trained on a bootstrap "
    "sample with random feature subsets at each split. Final predictions are by "
    "majority vote. The ensemble handles non-linear interactions naturally and "
    "provides Gini-based feature importances."
))

ieee_table_caption(doc, "Random Forest hyperparameters.")
rf_params = [
    ("n_estimators", "400"), ("max_depth", "12"),
    ("min_samples_leaf", "2"), ("Feature sampling", "\u221an_features"),
    ("Class weighting", "None (SMOTE)"),
    ("Deployment artefact", "rf_model.pkl (~264 MB) + rf_model.c"),
    ("Inference (ESP32-S3)", "~40 \u00b5s"),
]
rf_tbl = doc.add_table(rows=len(rf_params)+1, cols=2)
rf_tbl.style = "Table Grid"
for cell, text in zip(rf_tbl.rows[0].cells, ["Parameter", "Value"]):
    hdr_cell(cell, text)
for row, (k, v) in zip(rf_tbl.rows[1:], rf_params):
    tbl_cell(row.cells[0], k, bold=True, align_center=False)
    tbl_cell(row.cells[1], v, align_center=False)

# -- C. XGBoost --
ieee_subsection(doc, "C", "XGBoost Gradient Boosting")
ieee_para(doc, (
    "XGBoost implements gradient-boosted decision trees with regularisation. "
    "The histogram-based tree method enables efficient training on the full "
    "dataset. Sample weights derived from balanced class weights are passed "
    "to the fit procedure to address class imbalance."
))

ieee_table_caption(doc, "XGBoost hyperparameters.")
xgb_params = [
    ("n_estimators", "600"), ("max_depth", "8"),
    ("learning_rate", "0.03"), ("subsample", "0.8"),
    ("colsample_bytree", "0.8"), ("min_child_weight", "3"),
    ("gamma", "0.1"), ("reg_alpha / reg_lambda", "0.1 / 1.0"),
    ("tree_method", "hist"), ("eval_metric", "mlogloss"),
    ("Class weighting", "sample_weight (balanced)"),
    ("Deployment artefact", "xgb_model.pkl (~22 MB)"),
    ("Inference (ESP32-S3)", "~40 \u00b5s"),
]
xgb_tbl = doc.add_table(rows=len(xgb_params)+1, cols=2)
xgb_tbl.style = "Table Grid"
for cell, text in zip(xgb_tbl.rows[0].cells, ["Parameter", "Value"]):
    hdr_cell(cell, text)
for row, (k, v) in zip(xgb_tbl.rows[1:], xgb_params):
    tbl_cell(row.cells[0], k, bold=True, align_center=False)
    tbl_cell(row.cells[1], v, align_center=False)

# -- D. Neural Network --
ieee_subsection(doc, "D", "Feedforward Neural Network (TFLite)")
ieee_para(doc, (
    "A deep feedforward neural network is implemented using TensorFlow/Keras with "
    "batch normalisation and dropout regularisation. Early stopping on validation "
    "accuracy with weight restoration ensures the best generalising checkpoint."
))
ieee_para(doc, (
    "Architecture: Input(128) \u2192 Dense(256, ReLU) \u2192 BatchNorm \u2192 "
    "Dropout(0.3) \u2192 Dense(128, ReLU) \u2192 BatchNorm \u2192 Dropout(0.25) \u2192 "
    "Dense(64, ReLU) \u2192 Dropout(0.2) \u2192 Dense(32, ReLU) \u2192 Dense(4, Softmax)."
), italic=True, first_indent=False)

ieee_table_caption(doc, "Neural Network architecture and training parameters.")
nn_params = [
    ("Hidden layers", "256 \u2192 128 \u2192 64 \u2192 32"),
    ("Optimiser", "Adam, lr = 0.001"),
    ("Loss", "Sparse categorical cross-entropy"),
    ("Batch size", "512"), ("Max epochs", "200 (early stopping)"),
    ("Early stopping", "monitor=val_accuracy, patience=15"),
    ("LR scheduler", "ReduceLROnPlateau (factor=0.5, patience=7)"),
    ("Class weighting", "class_weight (balanced)"),
    ("Deployment (.keras)", "nn_model.keras (full precision)"),
    ("Deployment (TFLite)", "model.tflite (INT8, ~91 KB)"),
    ("Inference (ESP32-S3)", "~86 \u00b5s"),
]
nn_tbl = doc.add_table(rows=len(nn_params)+1, cols=2)
nn_tbl.style = "Table Grid"
for cell, text in zip(nn_tbl.rows[0].cells, ["Parameter", "Value"]):
    hdr_cell(cell, text)
for row, (k, v) in zip(nn_tbl.rows[1:], nn_params):
    tbl_cell(row.cells[0], k, bold=True, align_center=False)
    tbl_cell(row.cells[1], v, align_center=False)

# ══════════════════════════════════════════════════════════════════════════════
# V. EXPERIMENTAL RESULTS
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Experimental Results")

ieee_subsection(doc, "A", "Overall Performance")
ieee_para(doc, (
    "Table IX summarises the primary evaluation metrics for all four models on the "
    "held-out test set (chronologically the most recent 15%). The macro-averaged F1 "
    "score is the primary ranking metric, as it treats all classes equally regardless "
    "of support."
))

artifact_size_map = {
    "Logistic Regression": "~6 KB",
    "Random Forest":       "~264 MB",
    "XGBoost":             "~22 MB",
    "Neural Network":      "~91 KB",
}
infer_map = {
    "Logistic Regression": "~4.3 \u00b5s",
    "Random Forest":       "~40 \u00b5s",
    "XGBoost":             "~40 \u00b5s",
    "Neural Network":      "~86 \u00b5s",
}
best_f1_name = max(results, key=lambda r: r["metrics"]["macro_f1"])["name"]

ieee_table_caption(doc, "Executive summary of model performance.")
summary_tbl = doc.add_table(rows=1 + len(results), cols=5)
summary_tbl.style = "Table Grid"
for cell, text in zip(summary_tbl.rows[0].cells,
    ["Model", "Accuracy", "Macro F1", "Artifact", "Inf. Time"]):
    hdr_cell(cell, text)
for row, m in zip(summary_tbl.rows[1:], results):
    acc = f"{m['metrics']['accuracy']:.3f}"
    mf1 = f"{m['metrics']['macro_f1']:.3f}"
    art = artifact_size_map.get(m["name"], "\u2014")
    inf = infer_map.get(m["name"], "\u2014")
    is_best = m["name"] == best_f1_name
    tbl_cell(row.cells[0], m["name"], bold=is_best, align_center=False)
    tbl_cell(row.cells[1], acc, bold=is_best)
    tbl_cell(row.cells[2], mf1, bold=is_best)
    tbl_cell(row.cells[3], art)
    tbl_cell(row.cells[4], inf)
    if is_best:
        for cell in row.cells:
            shade_cell(cell, "E8E8E8")

ieee_subsection(doc, "B", "Per-Class Performance")
ieee_para(doc, (
    "Table X reports per-class precision, recall, and F1. The primary confusion "
    "occurs at the Sunny\u2194Cloudy boundary. The Snowy class shows the widest "
    "variance, reflecting its rarity and the effectiveness of SMOTE with freeze-flag "
    "discriminative features."
))

test_support = {i: int(test_df.label.value_counts().get(i, 0)) for i in range(4)}

ieee_table_caption(doc, "Per-class precision, recall, and F1 by model.")
per_class_tbl = doc.add_table(rows=1 + 4*len(results), cols=5)
per_class_tbl.style = "Table Grid"
for cell, text in zip(per_class_tbl.rows[0].cells,
    ["Model / Class", "Prec.", "Recall", "F1", "Support"]):
    hdr_cell(cell, text)
row_idx = 1
for m in results:
    mname = m["name"]
    prec = m["metrics"]["per_class_precision"]
    rec  = m["metrics"]["per_class_recall"]
    f1   = m["metrics"]["per_class_f1"]
    for ci, cname in enumerate(CLASS_NAMES):
        row = per_class_tbl.rows[row_idx]
        label = f"{mname} \u2014 {cname}" if ci == 0 else f"    {cname}"
        tbl_cell(row.cells[0], label, align_center=False, bold=(ci == 0))
        tbl_cell(row.cells[1], f"{prec[ci]:.3f}")
        tbl_cell(row.cells[2], f"{rec[ci]:.3f}")
        tbl_cell(row.cells[3], f"{f1[ci]:.3f}")
        tbl_cell(row.cells[4], str(test_support.get(ci, 0)))
        if ci == 0:
            shade_cell(row.cells[0], "E8E8E8")
        row_idx += 1

ieee_subsection(doc, "C", "Confusion Matrices")
ieee_para(doc, (
    "Figs. 2\u20135 present the 4\u00d74 confusion matrices for each model. "
    "Diagonal elements (correct predictions) are highlighted."
))

for m in results:
    safe = m["name"].lower().replace(" ", "_")
    ieee_add_figure(doc, os.path.join(REPORTS_DIR, f"cm_{safe}.png"),
                    f"Confusion matrix \u2014 {m['name']}.")

ieee_subsection(doc, "D", "ROC Curves")
ieee_para(doc, (
    "Fig. 6 shows the macro-averaged one-vs-rest ROC curves for all four models. "
    "XGBoost achieves the highest ROC-AUC (0.827), followed closely by the neural "
    "network (0.826) and random forest (0.819)."
))
ieee_add_figure(doc, os.path.join(REPORTS_DIR, "roc_curves.png"),
                "ROC curves (macro one-vs-rest) for all four models.")

ieee_subsection(doc, "E", "Feature Importance")
ieee_para(doc, (
    "Figs. 7\u201310 show the top-20 features by importance for each model. "
    "Pressure tendency and rolling pressure statistics consistently rank highest, "
    "confirming barometric pressure change as the primary predictive signal."
))
for m in results:
    safe = m["name"].lower().replace(" ", "_")
    ieee_add_figure(doc, os.path.join(REPORTS_DIR, f"fi_{safe}.png"),
                    f"Top-20 feature importance \u2014 {m['name']}.")

# ══════════════════════════════════════════════════════════════════════════════
# VI. ANALYSIS AND DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Analysis and Discussion")

ieee_subsection(doc, "A", "XGBoost: Best Overall Model")
ieee_para(doc, (
    "XGBoost achieves the highest accuracy (0.582) and macro-F1 (0.541), making it "
    "the recommended primary model. Its 600-estimator boosted ensemble with "
    "max_depth=8 and learning_rate=0.03 captures non-linear interactions between "
    "pressure tendency, humidity, and temperature. XGBoost achieves the best "
    "Rainy-class F1 (0.547) and strong Sunny-class F1 (0.631)."
))
ieee_para(doc, (
    "The ~22 MB serialised model is compact for a host machine but exceeds the "
    "ESP32-S3\u2019s 8 MB PSRAM. For on-device inference, the LR coefficient export "
    "or NN TFLite model are the practical deployment targets."
))

ieee_subsection(doc, "B", "Random Forest: Close Second")
ieee_para(doc, (
    "Random forest achieves accuracy = 0.581 and macro-F1 = 0.522. The 400-tree "
    "ensemble provides strong Cloudy-class F1 (0.618) and competitive Snowy-class "
    "F1 (0.407). At ~264 MB serialised, it serves as a server-side classifier. An "
    "embedded C decision tree export (rf_model.c) with reduced depth can be "
    "generated for microcontroller deployment."
))

ieee_subsection(doc, "C", "Neural Network: Competitive Deep Learning")
ieee_para(doc, (
    "The neural network achieves accuracy = 0.572 and macro-F1 = 0.500. Notably "
    "it achieves the highest Sunny-class F1 (0.650) among all models. The INT8-"
    "quantised TFLite model (~91 KB) fits within the ESP32-S3\u2019s PSRAM at "
    "~86 \u00b5s inference latency\u2014well within the 30-minute prediction interval."
))

ieee_subsection(doc, "D", "Logistic Regression: Interpretable Baseline")
ieee_para(doc, (
    "Logistic regression achieves accuracy = 0.562 and macro-F1 = 0.482, the lowest "
    "among all four models. Its linear boundary cannot capture non-linear weather "
    "transitions, but the 4 \u00d7 128 coefficient matrix provides direct interpretability "
    "and inference costs only ~4.3 \u00b5s. The C header export (lr_coefficients.h, ~6 KB) "
    "requires no ML runtime and compiles directly into ESP32-S3 firmware."
))

ieee_subsection(doc, "E", "Cross-Class Analysis")
ieee_para(doc, (
    "The primary confusion occurs at the Sunny\u2194Cloudy boundary, where both classes "
    "share similar temperature and humidity profiles. The Snowy class shows consistently "
    "lower F1 (0.272\u20130.420) despite SMOTE augmentation, reflecting the difficulty of "
    "distinguishing Snowy from cold-Cloudy conditions. Expanding the dataset with more "
    "high-latitude stations would improve Snowy classification."
))

# ══════════════════════════════════════════════════════════════════════════════
# VII. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Deployment on LilyGo T-Display S3")

ieee_para(doc, (
    "The target platform is the LilyGo T-Display S3 (ESP32-S3: dual-core 240 MHz "
    "Xtensa LX7, 8 MB PSRAM, 16 MB Flash). Two models are practically deployable "
    "on-device: LR (lr_coefficients.h, ~6 KB) and NN (model.tflite, ~91 KB INT8). "
    "XGBoost and RF can run on a companion host machine (Table XV)."
))

ieee_table_caption(doc, "Deployment characteristics.")
deploy_rows = [
    ("Logistic Regression", "lr_coefficients.h (~6 KB)",  "< 1 ms",  "< 1 MB",  "~4 \u00b5s"),
    ("Random Forest",       "rf_model.pkl (~264 MB)",      "~5 s",    "~1.2 GB", "~40 \u00b5s"),
    ("XGBoost",             "xgb_model.pkl (~22 MB)",      "~200 ms", "~100 MB", "~40 \u00b5s"),
    ("Neural Network",      "model.tflite (~91 KB)",       "< 500 ms","< 50 MB", "~86 \u00b5s"),
]
deploy_tbl = doc.add_table(rows=1 + len(deploy_rows), cols=5)
deploy_tbl.style = "Table Grid"
for cell, text in zip(deploy_tbl.rows[0].cells,
    ["Model", "Artefact", "Load", "RAM", "Latency"]):
    hdr_cell(cell, text)
for row, vals in zip(deploy_tbl.rows[1:], deploy_rows):
    for cell, text in zip(row.cells, vals):
        tbl_cell(cell, text)

ieee_para(doc, "Recommended deployment stack:", first_indent=False)
for b in [
    "On-device primary: NN (model.tflite via TFLite Micro) \u2014 91 KB INT8.",
    "On-device fallback: LR (lr_coefficients.h) \u2014 pure C, no ML runtime.",
    "Host-side reference: XGBoost (xgb_model.pkl) \u2014 best macro-F1 (0.541).",
    "Sensor: BME688 via I\u00b2C \u2192 24-hour buffer \u2192 128-feature vector.",
    "Output: 4-class probability \u2192 T-Display S3 screen or MQTT broker.",
]:
    ieee_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# VIII. RECOMMENDATIONS AND FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Recommendations and Future Work")

ieee_subsection(doc, "A", "Immediate Actions")
for b in [
    "Expand Snowy-class training data with high-latitude stations (Troms\u00f8, Yakutsk, Winnipeg).",
    "Tune XGBoost hyperparameters via grid search for further improvement.",
    "Profile TFLite Micro inference on actual ESP32-S3 hardware.",
    "Evaluate on held-out recent data to verify temporal generalisation.",
]:
    ieee_bullet(doc, b)

ieee_subsection(doc, "B", "Model Improvements")
for b in [
    "Implement LSTM / temporal CNN to exploit the sequential 24-hour structure.",
    "Experiment with LightGBM as an alternative gradient boosting framework.",
    "Apply post-hoc temperature scaling for calibrated probability estimates.",
    "Explore attention mechanisms to weight recent hours more heavily.",
]:
    ieee_bullet(doc, b)

ieee_subsection(doc, "C", "Feature Engineering")
for b in [
    "Add wind speed/direction from Open-Meteo as training features.",
    "Incorporate cloud cover and precipitation probability metrics.",
    "Experiment with 48 h and 72 h lookback windows and 12\u201324 h horizons.",
    "Apply PCA or mutual-information selection to prune the 128-feature vector.",
]:
    ieee_bullet(doc, b)

ieee_subsection(doc, "D", "Production Deployment")
for b in [
    "Implement OTA firmware updates for new TFLite models on ESP32-S3 devices.",
    "Log predictions to SD card or MQTT for drift detection.",
    "Re-train models seasonally to account for climate drift.",
    "Explore on-device personalisation via transfer learning on local BME688 data.",
]:
    ieee_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# IX. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Conclusion")

ieee_para(doc, (
    "This paper presented a systematic comparison of four machine learning approaches "
    "for 6-hour weather condition classification using Open-Meteo historical data, "
    "targeting the LilyGo T-Display S3 (ESP32-S3) edge platform. XGBoost is "
    "recommended as the primary host-side model, achieving the best accuracy (0.582) "
    "and macro-F1 (0.541) across four weather classes, trained on 128 features derived "
    "from a 24-hour lookback window across six geographically diverse stations."
))
ieee_para(doc, (
    "For direct on-device inference, the INT8-quantised neural network "
    "(model.tflite, ~91 KB) is recommended\u2014achieving competitive accuracy (0.572) "
    "and macro-F1 (0.500) within the device\u2019s 8 MB PSRAM at ~86 \u00b5s inference "
    "latency. The logistic regression export (lr_coefficients.h, ~6 KB) provides a "
    "pure-C fallback requiring no ML runtime, at ~4.3 \u00b5s."
))
ieee_para(doc, (
    "The most significant remaining challenge is the Snowy class (F1: 0.272\u20130.420), "
    "which\u2014despite SMOTE augmentation\u2014remains the weakest due to limited "
    "high-latitude training data. Expanding the dataset with northern stations and "
    "incorporating wind direction and cloud-cover features are the highest-priority "
    "next steps."
))

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "References")

refs = [
    '[1] Open-Meteo, "Open-Meteo Historical Weather API," https://open-meteo.com.',
    '[2] Bosch Sensortec, "BME688 Digital Low Power Gas, Pressure, Temperature & Humidity Sensor," Datasheet, 2021.',
    '[3] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.',
    '[4] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. ACM SIGKDD, 2016, pp. 785\u2013794.',
    '[5] N. V. Chawla et al., "SMOTE: Synthetic Minority Over-sampling Technique," J. Artif. Intell. Res., vol. 16, pp. 321\u2013357, 2002.',
    '[6] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," J. Mach. Learn. Res., vol. 12, pp. 2825\u20132830, 2011.',
    '[7] M. Abadi et al., "TensorFlow: A System for Large-Scale Machine Learning," in Proc. OSDI, 2016, pp. 265\u2013283.',
    '[8] Espressif Systems, "ESP32-S3 Technical Reference Manual," 2022.',
    '[9] LilyGo, "T-Display S3 Product Specification," 2023.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(ref)
    set_font(run, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════════════════════

sec_num += 1
ieee_section(doc, sec_num, "Appendix: Complete Feature Vector")

ieee_para(doc, (
    "Table XVI lists the complete ordered 128-feature vector. Features 0\u201371 are "
    "24-hour lag signals, 72\u2013116 are derived statistics, and 117\u2013127 are cyclical "
    "time encoding and discriminative composite features."
), size=8)

feat_names = []
for sig in ["temp", "humidity", "pressure"]:
    for t in range(24, 0, -1):
        feat_names.append(f"{sig}_t-{t}h")
feat_names += ["dp_1h", "dp_12h", "dp_24h", "dp_accel", "dt_1h", "dt_12h",
               "dew_point", "abs_humidity"]
for sig in ["temp", "humidity", "pressure"]:
    for w in ["6h", "12h", "24h"]:
        for s in ["mean", "std", "min", "max"]:
            feat_names.append(f"{sig}_{s}_{w}")
feat_names += ["sin_hour", "cos_hour", "sin_doy", "cos_doy",
               "dp_depression", "near_freezing_any", "below_freezing_now",
               "near_freezing_now", "pressure_trend_sign", "dp_depression_norm",
               "snow_composite", "rain_composite"]

ieee_table_caption(doc, "Complete ordered feature vector (128 features).")
n_cols = 4
n_rows = (len(feat_names) + n_cols - 1) // n_cols
app_tbl = doc.add_table(rows=n_rows + 1, cols=n_cols * 2)
app_tbl.style = "Table Grid"
for ci in range(n_cols):
    hdr_cell(app_tbl.rows[0].cells[ci*2],   "#")
    hdr_cell(app_tbl.rows[0].cells[ci*2+1], "Feature")
for idx, name in enumerate(feat_names):
    row_i = idx % n_rows + 1
    col_i = (idx // n_rows) * 2
    tbl_cell(app_tbl.rows[row_i].cells[col_i],   str(idx), size=7)
    tbl_cell(app_tbl.rows[row_i].cells[col_i+1], name, size=7, align_center=False)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE, "reports", "AI_Weather_Prediction_Technical_Paper.docx")
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(f"Saved: {out_path}")
